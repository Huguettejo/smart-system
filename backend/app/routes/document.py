
from flask import Blueprint, request, jsonify
from flask_jwt_extended import jwt_required, get_jwt_identity
from ..models.document import Document
from ..models.user import Utilisateur
from ..extensions import db
from ..services.hugging_face_service import HuggingFaceService
import PyPDF2
import docx
import io
from werkzeug.utils import secure_filename


# CONFIGURATION UPLOAD
document_bp = Blueprint('document', __name__, url_prefix='/api/documents')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# EXTRACTION DE TEXTE DES FICHIERS
def extraire_texte_pdf(file_content):
    """Extraire le texte d'un fichier PDF"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        texte = ""
        for page in pdf_reader.pages:
            texte += page.extract_text() + "\n"
        return texte.strip()
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction du PDF: {str(e)}")

def extraire_texte_docx(file_content):
    """Extraire le texte d'un fichier DOCX"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        texte = ""
        for paragraph in doc.paragraphs:
            texte += paragraph.text + "\n"
        return texte.strip()
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction du DOCX: {str(e)}")


# ROUTE UPLOAD DOCUMENT

@document_bp.route('/upload', methods=['POST'])
@jwt_required()
def upload_document():
    """Upload et traitement d'un document pour génération de QCM"""
    try:
        current_user_id = get_jwt_identity()
        user = Utilisateur.query.get(current_user_id)
        if not user or user.role != 'enseignant':
            return jsonify({'error': 'Accès refusé. Seuls les enseignants peuvent uploader des documents.'}), 403

        if 'file' not in request.files:
            return jsonify({'error': 'Aucun fichier fourni'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Aucun fichier sélectionné'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Format de fichier non supporté. Utilisez PDF, TXT ou DOCX'}), 400

        file_content = file.read()
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()

        try:
            if file_extension == 'pdf':
                texte_contenu = extraire_texte_pdf(file_content)
            elif file_extension == 'docx':
                texte_contenu = extraire_texte_docx(file_content)
            elif file_extension == 'txt':
                texte_contenu = file_content.decode('utf-8')
            else:
                return jsonify({'error': 'Format non supporté'}), 400
        except Exception as e:
            return jsonify({'error': f"Erreur lors de l'extraction du texte: {str(e)}"}), 500

        # Sauvegarde du document en base
        document = Document(
            titre=filename,
            type=file_extension,
            contenu=texte_contenu,
            enseignant_id=current_user_id
        )
        document.sauvegarder()

        return jsonify({'message': 'Document uploadé avec succès', 'document': document.to_dict()}), 201

    except Exception as e:
        return jsonify({'error': f"Erreur lors de l'upload: {str(e)}"}), 500


# ROUTE GENERATION QUIZ

@document_bp.route('/<int:document_id>/generer-quiz', methods=['POST'])
@jwt_required()
def generer_quiz_depuis_document():
    """Générer un quiz à partir d'un document avec Hugging Face"""
    try:
        current_user_id = get_jwt_identity()
        document_id = request.view_args['document_id']

        data = request.get_json() or {}
        nombre_qcm = data.get('nombre_qcm', 50)
        nombre_vrai_faux = data.get('nombre_vrai_faux', 30)
        nombre_ouvertes = data.get('nombre_ouvertes', 20)

        document = Document.query.filter_by(id=document_id, enseignant_id=current_user_id).first()
        if not document:
            return jsonify({'error': 'Document non trouvé ou accès refusé'}), 404

        hf_service = HuggingFaceService()
        questions_qcm = hf_service.generer_questions_qcm(document.contenu, nombre_qcm)
        questions_vf = hf_service.generer_questions_vrai_faux(document.contenu, nombre_vrai_faux)
        questions_ouvertes = hf_service.generer_questions_ouvertes(document.contenu, nombre_ouvertes)

        questions_sauvegardees = []

        # Sauvegarde des questions
        for question_data in questions_qcm + questions_vf + questions_ouvertes:
            from ..models.qcm import QCM, OptionReponse
            qcm = QCM(
                titre=f"Question {len(questions_sauvegardees) + 1}",
                contenu=question_data['question'],
                type_exercice=question_data['type'],
                difficulte=question_data['difficulte'],
                reponse_attendue=str(question_data['reponse_correcte']),
                document_id=document_id
            )
            qcm.publier()

            if question_data['options']:
                for option_data in question_data['options']:
                    option = OptionReponse(
                        texte=option_data['texte'],
                        est_correcte=option_data['est_correcte'],
                        qcm_id=qcm.id
                    )
                    db.session.add(option)
                db.session.commit()

            questions_sauvegardees.append(qcm.to_dict())

        return jsonify({
            'message': f'{len(questions_sauvegardees)} questions générées avec succès',
            'total_questions': len(questions_sauvegardees),
            'breakdown': {
                'qcm': len(questions_qcm),
                'vrai_faux': len(questions_vf),
                'ouvertes': len(questions_ouvertes)
            },
            'document_id': document_id
        }), 201

    except Exception as e:
        return jsonify({'error': f'Erreur lors de la génération: {str(e)}'}), 500


# ROUTES DE LISTE ET DÉTAIL DOCUMENT
@document_bp.route('/', methods=['GET'])
@jwt_required()
def lister_documents():
    try:
        current_user_id = get_jwt_identity()
        user = Utilisateur.query.get(current_user_id)
        if user.role != 'enseignant':
            return jsonify({'error': 'Accès refusé'}), 403

        documents = Document.query.filter_by(enseignant_id=current_user_id).all()
        return jsonify({'documents': [doc.to_dict() for doc in documents]}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@document_bp.route('/<int:document_id>', methods=['GET'])
@jwt_required()
def obtenir_document():
    try:
        current_user_id = get_jwt_identity()
        document_id = request.view_args['document_id']

        document = Document.query.filter_by(id=document_id, enseignant_id=current_user_id).first()
        if not document:
            return jsonify({'error': 'Document non trouvé'}), 404

        return jsonify({'document': document.to_dict()}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ROUTE SUPPRESSION DOCUMENT
@document_bp.route('/<int:document_id>', methods=['DELETE'])
@jwt_required()
def supprimer_document():
    try:
        current_user_id = get_jwt_identity()
        document_id = request.view_args['document_id']

        document = Document.query.filter_by(id=document_id, enseignant_id=current_user_id).first()
        if not document:
            return jsonify({'error': 'Document non trouvé'}), 404

        document.supprimer()
        return jsonify({'message': 'Document supprimé avec succès'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
